import docx, pickle
from tkinter import *

File = "FTH Notes/FTH Ch1.note"

def to_docx(file):
    with open(file,'rb') as f:
        data = pickle.load(f)
    name, subject, date, keywords, notes, summary = data

    doc = docx.Document()
    doc.add_heading(name,0)
    doc.add_heading(subject,0)
    doc.add_heading(date,0)
    doc.add_paragraph(keywords)
    doc.add_paragraph(notes)
    doc.add_paragraph(summary)
    
    doc.save((file.replace(".note","")) + ".docx")

to_docx(File)
